"""Verify (a) docx download contains generated section text, (b) student role gating."""
import io
import json
import os
from pathlib import Path

import requests
from docx import Document
from dotenv import dotenv_values

BASE_URL = (os.environ.get("REACT_APP_BACKEND_URL")
            or dotenv_values("/app/frontend/.env")["REACT_APP_BACKEND_URL"]).rstrip("/")
fx = json.loads(Path("/app/backend/tests/.lesson_plan_fixture.json").read_text())

tok = requests.post(f"{BASE_URL}/api/auth/teacher-login",
                    json={"email": "astapp@spanola.net", "password": "AlisaFaith$14"},
                    timeout=60).json()["session_token"]
h = {"Authorization": f"Bearer {tok}"}

plans = requests.get(f"{BASE_URL}/api/lesson-plans", headers=h, timeout=90).json()
target = next((p for p in plans if p.get("template_id") == fx["template_id"]), None)
print("target plan:", target and target["id"], "| title:", target and target.get("title"))
sections = target["plans"][0]["sections"]

r = requests.get(f"{BASE_URL}/api/lesson-plans/{target['id']}/download/0", headers=h, timeout=120)
print("docx download HTTP", r.status_code, "bytes", len(r.content),
      "| content-type", r.headers.get("content-type"))
if r.status_code == 200 and r.content[:2] == b"PK":
    doc = Document(io.BytesIO(r.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    print("docx chars:", len(text))
    for lbl in ("Learner Outcomes", "Standards", "Closure"):
        snippet = (sections.get(lbl) or "")[:60]
        print(f"  {lbl!r}: label_present={lbl in text} content_present={bool(snippet) and snippet in text}")
    print("---- docx text preview ----")
    print(text[:1200])
else:
    print(r.text[:400])

# Student role gating
print("\n=== student gating ===")
for path, body in [("/api/auth/login", {"username": "student", "password": "wrong"})]:
    rr = requests.post(f"{BASE_URL}{path}", json=body, timeout=60)
    print(path, rr.status_code, rr.text[:150])
