"""Fixture setup: insert a TEST_ fillable docx template + find a turtle lesson with problems."""
import asyncio
import base64
import io
import json
import os
import uuid

from dotenv import dotenv_values
from motor.motor_asyncio import AsyncIOMotorClient
from docx import Document

env = dotenv_values("/app/backend/.env")
MONGO_URL = env.get("MONGO_URL") or os.environ["MONGO_URL"]
DB_NAME = env.get("DB_NAME") or os.environ["DB_NAME"]

LABELS = [
    "Learner Outcomes", "Standards", "Anticipatory Set", "Modeling",
    "Instructional Strategies", "Check for Understanding", "Guided Practice",
    "Independent Practice", "Closure", "Differentiation", "Assessment",
]


def build_docx_b64():
    doc = Document()
    doc.add_paragraph("TEST_ Lesson Plan Template")
    for label in LABELS:
        doc.add_paragraph(f"{label}:")
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode()


async def main():
    client = AsyncIOMotorClient(MONGO_URL)
    db = client[DB_NAME]

    user = await db.users.find_one({"email": "astapp@spanola.net"})
    assert user, "admin teacher not found"
    user_id = user["id"]

    # Find a turtle lesson that has problems
    pipeline = [
        {"$match": {"assignment_type": "turtle"}},
        {"$group": {"_id": {"chapter": "$chapter", "lesson": "$lesson"}, "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
        {"$limit": 5},
    ]
    combos = [c async for c in db.problems.aggregate(pipeline)]
    print("turtle combos:", combos)

    tpl_id = "TEST_tpl_" + str(uuid.uuid4())[:8]
    await db.lesson_plan_templates.delete_many({"name": {"$regex": "^TEST_"}})
    await db.lesson_plan_templates.insert_one({
        "id": tpl_id,
        "user_id": user_id,
        "name": "TEST_Fillable_Template.docx",
        "format": "docx",
        "fillable": True,
        "section_labels": LABELS,
        "extracted_text": "\n".join(f"{lbl}:\n" for lbl in LABELS),
        "file_b64": build_docx_b64(),
        "uploaded_at": "2026-07-01T00:00:00+00:00",
    })

    out = {
        "template_id": tpl_id,
        "user_id": user_id,
        "combos": [{"chapter": c["_id"]["chapter"], "lesson": c["_id"]["lesson"], "n": c["n"]} for c in combos],
    }
    with open("/app/backend/tests/.lesson_plan_fixture.json", "w") as f:
        json.dump(out, f, indent=2)
    print(json.dumps(out, indent=2))
    client.close()


asyncio.run(main())
