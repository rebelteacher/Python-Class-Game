"""
Tests for the AI Lesson Plan Generator schedule + templates endpoints.
Covers the two-bug fix session:
  1) generate-from-schedule respects assignment_type from the frontend
     (turtle unit produces turtle-content, NOT Python print()).
  2) generate-from-schedule refuses to hallucinate when there is no
     grounding data — returns a "No content found" plan with the ⚠️ marker.
  3) seed_turtle_problems and seed_microbit_problems are non-destructive.
  4) Template upload / list / delete flow works for .docx.
"""
import io
import os
import re
import pytest
import requests

def _load_backend_url():
    v = os.environ.get("REACT_APP_BACKEND_URL", "").strip()
    if not v:
        try:
            with open("/app/frontend/.env") as f:
                for line in f:
                    if line.startswith("REACT_APP_BACKEND_URL="):
                        v = line.split("=", 1)[1].strip().strip('"').strip("'")
                        break
        except FileNotFoundError:
            pass
    return v.rstrip("/")


BASE_URL = _load_backend_url()
ADMIN_EMAIL = "astapp@spanola.net"
ADMIN_PASSWORD = "AlisaFaith$14"


@pytest.fixture(scope="module")
def auth_headers():
    r = requests.post(
        f"{BASE_URL}/api/auth/teacher-login",
        json={"email": ADMIN_EMAIL, "password": ADMIN_PASSWORD},
        timeout=30,
    )
    assert r.status_code == 200, f"login failed: {r.status_code} {r.text}"
    tok = r.json()["session_token"]
    return {"Authorization": f"Bearer {tok}", "Content-Type": "application/json"}


@pytest.fixture(scope="module")
def turtle_lesson(auth_headers):
    """Discover a real turtle lesson (unit/chapter/lesson) from the preview DB."""
    r = requests.get(f"{BASE_URL}/api/curriculum/units", headers=auth_headers, timeout=30)
    assert r.status_code == 200
    units = r.json()
    turtle_unit = next(u for u in units if u.get("assignment_type") == "turtle")
    unit_name = turtle_unit["name"]
    # pick first chapter+lesson that has problems
    for ch in turtle_unit["chapters"]:
        for l in ch["lessons"]:
            if l.get("problem_count", 0) > 0 and not l.get("is_orphan"):
                return {"unit": unit_name, "chapter": ch["name"], "lesson": l["name"]}
    pytest.skip("No turtle lesson with problems found")


# ---------------- Seed non-destructiveness ----------------
class TestSeedNonDestructive:
    def test_seed_functions_have_early_return(self):
        with open("/app/backend/server.py") as f:
            src = f.read()

        # Verify both seed functions include the "skipping seed (non-destructive)" log
        assert "Turtle problems exist" in src and "skipping seed (non-destructive)" in src
        assert "Micro:bit problems exist" in src

        # Verify no ACTIVE delete_many call in seed function bodies. Docstrings and
        # `#` comments that mention delete_many are fine — only executable code counts.
        for fn_name in ["seed_turtle_problems", "seed_microbit_problems"]:
            m = re.search(rf"async def {fn_name}\(.*?\).*?:\n(.*?)\nasync def ", src, re.DOTALL)
            assert m, f"could not locate {fn_name}"
            body = m.group(1)
            # Strip triple-quoted docstrings and `#` line comments
            body_clean = re.sub(r'"""[\s\S]*?"""', "", body)
            body_clean = re.sub(r"#[^\n]*", "", body_clean)
            assert "delete_many" not in body_clean, (
                f"{fn_name} still contains an active delete_many call — seed is destructive!"
            )

    def test_startup_log_reports_skipping_seed(self):
        import glob
        logs = ""
        for lp in glob.glob("/var/log/supervisor/backend.*.log"):
            with open(lp, errors="ignore") as f:
                logs += f.read()
        assert "Turtle problems exist" in logs and "skipping seed (non-destructive)" in logs
        assert "Micro:bit problems exist" in logs


# ---------------- Templates CRUD ----------------
class TestLessonPlanTemplates:
    def _make_docx_buffer(self):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        doc = Document()
        doc.add_heading("Anticipatory Set", level=2)
        doc.add_paragraph("Hook students at the start.")
        doc.add_heading("Modeling", level=2)
        doc.add_paragraph("Teacher demonstrates.")
        doc.add_heading("Guided Practice", level=2)
        doc.add_paragraph("Together with the teacher.")
        doc.add_heading("Independent Practice", level=2)
        doc.add_paragraph("Students work alone.")
        doc.add_heading("Closure", level=2)
        doc.add_paragraph("Wrap up the lesson.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def test_upload_list_delete_template(self, auth_headers):
        buf = self._make_docx_buffer()
        # Strip Content-Type header — requests will set the multipart boundary
        h = {k: v for k, v in auth_headers.items() if k.lower() != "content-type"}

        files = {"file": ("test_template.docx", buf,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"name": "TEST_template_schedule"}
        r = requests.post(f"{BASE_URL}/api/lesson-plans/templates",
                          headers=h, files=files, data=data, timeout=30)
        assert r.status_code == 200, f"upload failed: {r.status_code} {r.text}"
        body = r.json()
        assert "id" in body
        assert body["format"] == "docx"
        preview = body.get("preview", "")
        # All five section names should appear in the extracted preview
        for sect in ["Anticipatory Set", "Modeling", "Guided Practice",
                     "Independent Practice", "Closure"]:
            assert sect in preview, f"Missing section '{sect}' in preview"
        tid = body["id"]

        # LIST
        r2 = requests.get(f"{BASE_URL}/api/lesson-plans/templates",
                          headers=auth_headers, timeout=30)
        assert r2.status_code == 200
        assert any(t["id"] == tid for t in r2.json())

        # DELETE
        r3 = requests.delete(f"{BASE_URL}/api/lesson-plans/templates/{tid}",
                             headers=auth_headers, timeout=30)
        assert r3.status_code == 200
        assert r3.json().get("success") is True

        # Confirm gone
        r4 = requests.get(f"{BASE_URL}/api/lesson-plans/templates",
                          headers=auth_headers, timeout=30)
        assert not any(t["id"] == tid for t in r4.json())


# ---------------- generate-from-schedule ----------------
class TestGenerateFromSchedule:
    def test_turtle_lesson_produces_turtle_content(self, auth_headers, turtle_lesson):
        """The critical bug: turtle unit was producing Python print() lessons."""
        payload = {
            "schedule": [
                {
                    "day_label": "Monday, Jan 12",
                    "day_index": 0,
                    "unit": turtle_lesson["unit"],
                    "chapter": turtle_lesson["chapter"],
                    "lesson": turtle_lesson["lesson"],
                    "assignment_type": "turtle",
                    "span_days": 1,
                    "day_within_span": 1,
                }
            ],
            "header_fields": {
                "schoolName": "Test School",
                "teacherName": "Test Teacher",
                "className": "7th Grade CS",
                "timePerPeriod": "50",
            },
            "ai_generate_standards": False,
            "ai_generate_objectives": False,
        }
        r = requests.post(f"{BASE_URL}/api/lesson-plans/generate-from-schedule",
                          headers=auth_headers, json=payload, timeout=180)
        assert r.status_code == 200, f"generate failed: {r.status_code} {r.text[:500]}"
        data = r.json()
        assert "plans" in data or isinstance(data, list) or "days" in data or "results" in data, \
            f"Unexpected shape: {list(data.keys()) if isinstance(data, dict) else type(data)}"

        # Try to pull the plans list out of common shapes
        plans = None
        if isinstance(data, list):
            plans = data
        else:
            for k in ("plans", "days", "results", "generated"):
                if k in data:
                    plans = data[k]
                    break
        assert plans and len(plans) == 1, f"Expected 1 plan; got: {data}"
        content = plans[0].get("content") or plans[0].get("plan") or ""
        assert isinstance(content, str) and len(content) > 100, \
            f"Content too short/empty: {content[:200]}"

        low = content.lower()
        print(f"\n----- GENERATED CONTENT (first 800 chars) -----\n{content[:800]}\n---------")
        # Must reference turtle
        assert "turtle" in low, f"'turtle' missing in content. Preview:\n{content[:500]}"
        # Must reference at least one turtle command family
        assert re.search(r"\b(forward|backward|left|right|pen(up|down)?|move)\b", low), \
            f"No turtle commands present. Preview:\n{content[:500]}"
        # Must NOT be dominated by print() lesson
        # Allow rare mentions but not primary topic
        print_count = low.count("print(")
        assert "utilize the print()" not in low
        assert "print command" not in low
        # print() may legitimately appear once or twice (e.g. debugging), but should
        # not appear many times as the main topic.
        assert print_count < 5, (
            f"Content mentions print() {print_count} times — likely defaulted to Python text lesson."
        )

    def test_missing_lesson_returns_no_content_marker(self, auth_headers):
        """When there are no problems + no instructions, backend returns ⚠️ plan."""
        payload = {
            "schedule": [
                {
                    "day_label": "Tuesday, Jan 13",
                    "day_index": 1,
                    "unit": "Unit 2: Turtle Graphics",
                    "chapter": "TEST_Fake Chapter_ZZZ",
                    "lesson": "TEST_Fake Lesson_ZZZ",
                    "assignment_type": "turtle",
                    "span_days": 1,
                    "day_within_span": 1,
                }
            ],
            "header_fields": {"schoolName": "Test", "teacherName": "T",
                              "className": "C", "timePerPeriod": "50"},
        }
        r = requests.post(f"{BASE_URL}/api/lesson-plans/generate-from-schedule",
                          headers=auth_headers, json=payload, timeout=60)
        assert r.status_code == 200, f"{r.status_code} {r.text[:300]}"
        data = r.json()
        plans = data if isinstance(data, list) else (
            data.get("plans") or data.get("days") or data.get("results") or data.get("generated")
        )
        assert plans and len(plans) == 1
        content = plans[0].get("content", "")
        assert "⚠️" in content or "No content found" in content, \
            f"Expected refusal marker, got: {content[:400]}"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
