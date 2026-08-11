"""
Backend tests for the new template-fillable lesson plan generator + downloads.

Covers:
  1) POST /api/lesson-plans/generate-from-schedule with a Batesville-style
     template returns plans[0].sections as a NON-EMPTY dict keyed by the
     template's exact labels, with turtle-related content per section.
  2) GET /api/lesson-plans/{plan_id}/download/0 returns a valid .docx with
     the template's table rows filled in the second cell (not blank / not
     'Not specified').
  3) GET /api/lesson-plans/{plan_id}/download-week for a 3-day plan returns
     a single .docx with 3 filled table copies separated by page breaks.
"""
import io
import os
import re
import pytest
import requests


def _load_backend_url():
    v = os.environ.get("REACT_APP_BACKEND_URL", "").strip()
    if not v:
        try:
            with open("/app/frontend/.env") as f:
                for line in f:
                    if line.startswith("REACT_APP_BACKEND_URL="):
                        v = line.split("=", 1)[1].strip().strip('"').strip("'")
                        break
        except FileNotFoundError:
            pass
    return v.rstrip("/")


BASE_URL = _load_backend_url()
ADMIN_EMAIL = "astapp@spanola.net"
ADMIN_PASSWORD = "AlisaFaith$14"

# The Batesville-style labels the frontend/user cares about
BATESVILLE_LABELS = [
    "Teacher Name",
    "Learner Outcomes / Objectives",
    "Standards",
    "Materials Needed",
    "Anticipatory Set",
    "Modeling",
    "Guided Practice / Monitoring",
    "Independent Practice",
]


@pytest.fixture(scope="module")
def auth_headers():
    r = requests.post(
        f"{BASE_URL}/api/auth/teacher-login",
        json={"email": ADMIN_EMAIL, "password": ADMIN_PASSWORD},
        timeout=30,
    )
    assert r.status_code == 200, f"login failed: {r.status_code} {r.text}"
    tok = r.json()["session_token"]
    return {"Authorization": f"Bearer {tok}"}


def _make_batesville_docx_bytes():
    """Build an in-memory .docx template that looks like the Batesville lesson
    plan: a single 2-column table where col1 = label, col2 = blank fill target."""
    from docx import Document
    doc = Document()
    doc.add_heading("Lesson Plan", level=1)
    table = doc.add_table(rows=len(BATESVILLE_LABELS), cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(BATESVILLE_LABELS):
        table.rows[i].cells[0].text = label
        # second cell intentionally left blank so backend can fill it
        table.rows[i].cells[1].text = ""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@pytest.fixture(scope="module")
def uploaded_template(auth_headers):
    buf = _make_batesville_docx_bytes()
    files = {"file": ("batesville_template.docx", buf,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    data = {"name": "TEST_batesville_template"}
    r = requests.post(f"{BASE_URL}/api/lesson-plans/templates",
                      headers=auth_headers, files=files, data=data, timeout=30)
    assert r.status_code == 200, f"upload failed: {r.status_code} {r.text}"
    body = r.json()
    assert body.get("fillable") is True, f"Template not fillable — labels detected: {body.get('section_labels')}"
    detected = body.get("section_labels") or []
    # All Batesville labels should have been detected
    for lab in BATESVILLE_LABELS:
        assert lab in detected, f"Label {lab!r} not detected in {detected}"
    tid = body["id"]
    yield {"id": tid, "labels": detected}
    # cleanup
    requests.delete(f"{BASE_URL}/api/lesson-plans/templates/{tid}",
                    headers=auth_headers, timeout=30)


@pytest.fixture(scope="module")
def turtle_lesson(auth_headers):
    """Locate 'Unit 2 Turtle / Chapter 1: First Steps / Lesson 1: Name Your
    Turtle' (or fall back to any turtle lesson with problems)."""
    r = requests.get(f"{BASE_URL}/api/curriculum/units", headers=auth_headers, timeout=30)
    assert r.status_code == 200
    units = r.json()
    turtle_unit = next((u for u in units if u.get("assignment_type") == "turtle"), None)
    assert turtle_unit, "No turtle unit found"
    # Prefer the specific Chapter 1: First Steps / Lesson 1: Name Your Turtle
    for ch in turtle_unit["chapters"]:
        if "First Steps" in ch["name"]:
            for l in ch["lessons"]:
                if "Name Your Turtle" in l["name"] and l.get("problem_count", 0) > 0:
                    return {
                        "unit": turtle_unit["name"],
                        "chapter": ch["name"],
                        "lesson": l["name"],
                    }
    # Fallback
    for ch in turtle_unit["chapters"]:
        for l in ch["lessons"]:
            if l.get("problem_count", 0) > 0 and not l.get("is_orphan"):
                return {"unit": turtle_unit["name"], "chapter": ch["name"], "lesson": l["name"]}
    pytest.skip("No turtle lesson with problems found")


# ------------ Test 1: generate-from-schedule returns sections dict --------
class TestGenerateSectionsDict:
    def test_generate_produces_template_keyed_sections(self, auth_headers, uploaded_template, turtle_lesson):
        payload = {
            "template_id": uploaded_template["id"],
            "schedule": [{
                "day_label": "Monday, Jan 12",
                "day_index": 0,
                "unit": turtle_lesson["unit"],
                "chapter": turtle_lesson["chapter"],
                "lesson": turtle_lesson["lesson"],
                "assignment_type": "turtle",
                "span_days": 1,
                "day_within_span": 1,
            }],
            "header_fields": {
                "schoolName": "Test School",
                "teacherName": "Ms. Test",
                "className": "7th Grade CS",
                "timePerPeriod": "50",
                "lessonRange": "Jan 12",
            },
        }
        r = requests.post(f"{BASE_URL}/api/lesson-plans/generate-from-schedule",
                          headers={**auth_headers, "Content-Type": "application/json"},
                          json=payload, timeout=180)
        assert r.status_code == 200, f"{r.status_code} {r.text[:500]}"
        data = r.json()
        assert "id" in data and "plans" in data
        plans = data["plans"]
        assert len(plans) == 1
        p0 = plans[0]

        # NEW: sections dict must exist, be a dict, and NOT be empty
        sections = p0.get("sections")
        assert isinstance(sections, dict), f"sections not a dict: {type(sections)}"
        assert len(sections) > 0, f"sections dict is empty. content preview: {p0.get('content','')[:300]}"

        # At least half the template labels should be present as keys
        template_labels = uploaded_template["labels"]
        matched = [lab for lab in template_labels if lab in sections]
        assert len(matched) >= max(3, len(template_labels) // 2), (
            f"Only {len(matched)}/{len(template_labels)} template labels present in sections. "
            f"Got keys: {list(sections.keys())}"
        )

        # Each section value must be a non-empty string
        for lab, val in sections.items():
            assert isinstance(val, str), f"section {lab!r} value not str: {type(val)}"
            assert val.strip(), f"section {lab!r} is empty/whitespace"
            assert "not specified" not in val.lower(), (
                f"section {lab!r} contains 'Not specified' — AI didn't fill it in properly."
            )

        # At least one content section should mention turtle-y content
        all_text = "\n".join(sections.values()).lower()
        assert "turtle" in all_text, f"'turtle' missing from all sections. First 400: {all_text[:400]}"
        assert re.search(r"\b(forward|backward|left|right|pen(up|down)?|circle|move)\b", all_text), (
            f"No turtle-graphics commands in any section. Preview: {all_text[:400]}"
        )
        # Should NOT be dominated by print() lesson
        print_count = all_text.count("print(")
        assert print_count < 5, f"print() appeared {print_count}× — likely defaulted to Python text lesson"

        # Save plan_id + template_id on the class for downstream tests
        pytest.PLAN_ID_SINGLE = data["id"]


# ------------ Test 2: download/{day_index} returns filled docx -----------
class TestDownloadDayDocx:
    def test_download_day_returns_filled_docx(self, auth_headers, uploaded_template, turtle_lesson):
        # Ensure we have a plan (regenerate independently — don't rely on cross-class state)
        payload = {
            "template_id": uploaded_template["id"],
            "schedule": [{
                "day_label": "Monday, Jan 12",
                "day_index": 0,
                "unit": turtle_lesson["unit"],
                "chapter": turtle_lesson["chapter"],
                "lesson": turtle_lesson["lesson"],
                "assignment_type": "turtle",
                "span_days": 1,
                "day_within_span": 1,
            }],
            "header_fields": {
                "schoolName": "Test School",
                "teacherName": "Ms. Test",
                "className": "7th Grade CS",
                "timePerPeriod": "50",
                "lessonRange": "Jan 12",
                "standards": "CSTA 2-AP-13",
                "objectives": "SWBAT use turtle.forward()",
            },
            "ai_generate_standards": True,
            "ai_generate_objectives": True,
        }
        r = requests.post(f"{BASE_URL}/api/lesson-plans/generate-from-schedule",
                          headers={**auth_headers, "Content-Type": "application/json"},
                          json=payload, timeout=180)
        assert r.status_code == 200
        plan_id = r.json()["id"]

        # Download day 0
        rd = requests.get(f"{BASE_URL}/api/lesson-plans/{plan_id}/download/0",
                          headers=auth_headers, timeout=60)
        assert rd.status_code == 200, f"download failed: {rd.status_code} {rd.text[:300]}"
        ct = rd.headers.get("content-type", "")
        assert "wordprocessingml.document" in ct or "octet-stream" in ct, f"unexpected content-type: {ct}"
        assert len(rd.content) > 2000, f"docx too small: {len(rd.content)} bytes"

        from docx import Document
        doc = Document(io.BytesIO(rd.content))
        assert len(doc.tables) >= 1, "No tables in filled docx"

        # Check the first table — each row's first cell should be a known label
        # and the second cell should have non-empty, non-'Not specified' content.
        rows_with_content = 0
        labels_seen = []
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                label = (cells[0].text or "").strip()
                second = (cells[1].text or "").strip()
                if not label:
                    continue
                labels_seen.append(label)
                # Only assert on rows whose label we recognize
                if label in BATESVILLE_LABELS and label != "Teacher Name":
                    # Content section — must be filled (AI may skip a label if
                    # instructed there's no data, but with ai_generate=True most
                    # should have content).
                    if not second:
                        # Non-fatal — record it so we can assert on count below
                        continue
                    assert "not specified" not in second.lower(), (
                        f"Row {label!r} still contains 'Not specified': {second[:200]}"
                    )
                    rows_with_content += 1
        assert rows_with_content >= 5, (
            f"Only {rows_with_content} content rows filled in (need >=5). Labels seen: {labels_seen}"
        )


# ------------ Test 3: download-week returns multi-page docx --------------
class TestDownloadWeekDocx:
    def test_download_week_has_three_filled_tables_with_page_breaks(
            self, auth_headers, uploaded_template, turtle_lesson):
        payload = {
            "template_id": uploaded_template["id"],
            "schedule": [
                {"day_label": f"Day {i+1}", "day_index": i,
                 "unit": turtle_lesson["unit"],
                 "chapter": turtle_lesson["chapter"],
                 "lesson": turtle_lesson["lesson"],
                 "assignment_type": "turtle",
                 "span_days": 3,
                 "day_within_span": i + 1}
                for i in range(3)
            ],
            "header_fields": {
                "schoolName": "Test",
                "teacherName": "T",
                "className": "C",
                "timePerPeriod": "50",
                "lessonRange": "Jan 12-14",
            },
        }
        r = requests.post(f"{BASE_URL}/api/lesson-plans/generate-from-schedule",
                          headers={**auth_headers, "Content-Type": "application/json"},
                          json=payload, timeout=240)
        assert r.status_code == 200, f"{r.status_code} {r.text[:300]}"
        j = r.json()
        assert len(j["plans"]) == 3
        # sanity: each day should have a non-empty sections dict
        for i, dp in enumerate(j["plans"]):
            assert isinstance(dp.get("sections"), dict) and dp["sections"], (
                f"Day {i} has empty/missing sections dict"
            )
        plan_id = j["id"]

        rd = requests.get(f"{BASE_URL}/api/lesson-plans/{plan_id}/download-week",
                          headers=auth_headers, timeout=120)
        assert rd.status_code == 200, f"week download failed: {rd.status_code} {rd.text[:300]}"
        assert len(rd.content) > 3000, f"combined docx suspiciously small: {len(rd.content)}"

        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(io.BytesIO(rd.content))
        assert len(doc.tables) == 3, f"Expected 3 tables (one per day), got {len(doc.tables)}"

        # Count page breaks in the doc body
        body_xml = doc.element.body
        page_breaks = 0
        for br in body_xml.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                page_breaks += 1
        assert page_breaks >= 2, f"Expected >=2 page breaks separating 3 days, got {page_breaks}"

        # Each table should have Learner Outcomes filled in
        for i, table in enumerate(doc.tables):
            filled_row_count = 0
            for row in table.rows:
                cells = row.cells
                if len(cells) < 2:
                    continue
                label = (cells[0].text or "").strip()
                second = (cells[1].text or "").strip()
                if label in BATESVILLE_LABELS and label != "Teacher Name" and second:
                    assert "not specified" not in second.lower(), (
                        f"Day {i} row {label!r} says 'Not specified'"
                    )
                    filled_row_count += 1
            assert filled_row_count >= 3, f"Table {i} only has {filled_row_count} filled rows"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
